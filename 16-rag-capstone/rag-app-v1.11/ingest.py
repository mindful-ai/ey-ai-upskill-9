
from docx import Document
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.config import Settings

doc = Document("medical_knowledge_dataset.docx")
print(f"Total paragraphs in document: {len(doc.paragraphs)}")
print("Sample paragraph:", doc.paragraphs[0].text[:200])  # Print the first 200 characters of the first paragraph

# Filter out paragraphs that are too short (e.g., less than 50 characters)
paragraphs = [p.text for p in doc.paragraphs if len(p.text.strip()) > 50]
print(f"Total paragraphs after filtering: {len(paragraphs)}")
print("Sample filtered paragraph:", paragraphs[0][:200])  # Print the first 200

model = SentenceTransformer("all-MiniLM-L6-v2")

# client = chromadb.Client(Settings(persist_directory="./chroma_db"))
client = chromadb.PersistentClient(path="./chroma_db")
collection = client.get_or_create_collection("medical_docs")

embeddings = model.encode(paragraphs).tolist()

# ChromaDB has a max batch size limit; process in chunks
BATCH_SIZE = 2000
total_docs = len(paragraphs)

for i in range(0, total_docs, BATCH_SIZE):
    batch_end = min(i + BATCH_SIZE, total_docs)
    batch_docs = paragraphs[i:batch_end]
    batch_embeddings = embeddings[i:batch_end]
    batch_ids = [str(j) for j in range(i, batch_end)]
    
    collection.add(
        documents=batch_docs,
        embeddings=batch_embeddings,
        ids=batch_ids
    )
    print(f"Ingested batch {i // BATCH_SIZE + 1}: documents {i} to {batch_end}")

print(f"Medical dataset ingested into ChromaDB ({total_docs} total documents)")
