# pip install python-docx

from docx import Document
import chromadb
from chromadb.utils import embedding_functions

# ---------------- CONFIG ---------------- #
DOC_PATH = "medical_knowledge_dataset.docx"
CHROMA_PATH = "./chroma_db"
COLLECTION_NAME = "medical_docs"

MIN_CHAR_LENGTH = 50
CHUNK_SIZE = 500     # characters per chunk
CHUNK_OVERLAP = 100  # overlap between chunks
BATCH_SIZE = 500

# ---------------- LOAD DOCUMENT ---------------- #
doc = Document(DOC_PATH)

print(f"Total paragraphs in document: {len(doc.paragraphs)}")

# Filter useful paragraphs
paragraphs = [
    p.text.strip()
    for p in doc.paragraphs
    if len(p.text.strip()) > MIN_CHAR_LENGTH
]

print(f"Paragraphs after filtering: {len(paragraphs)}")

# ---------------- CHUNKING FUNCTION ---------------- #
def chunk_text(text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    chunks = []
    start = 0
    text_length = len(text)

    while start < text_length:
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)
        start += chunk_size - overlap

    return chunks

# Apply chunking
documents = []
for para in paragraphs:
    documents.extend(chunk_text(para))

print(f"Total chunks created: {len(documents)}")

# ---------------- CHROMA SETUP ---------------- #
embedding_function = embedding_functions.SentenceTransformerEmbeddingFunction(
    model_name="all-MiniLM-L6-v2"
)

client = chromadb.PersistentClient(path=CHROMA_PATH)

collection = client.get_or_create_collection(
    name=COLLECTION_NAME,
    embedding_function=embedding_function
)

# ---------------- AVOID DUPLICATE INGEST ---------------- #
if collection.count() > 0:
    print("⚠️ Collection already contains data. Skipping ingestion.")
    print(f"Existing document count: {collection.count()}")
    exit()

# ---------------- INGEST DATA ---------------- #
total_docs = len(documents)

for i in range(0, total_docs, BATCH_SIZE):
    batch_end = min(i + BATCH_SIZE, total_docs)

    batch_docs = documents[i:batch_end]
    batch_ids = [f"doc_{j}" for j in range(i, batch_end)]

    # Optional metadata
    batch_metadata = [
        {"source": "medical_doc", "chunk_id": j}
        for j in range(i, batch_end)
    ]

    collection.add(
        documents=batch_docs,
        ids=batch_ids,
        metadatas=batch_metadata
    )

    print(f"Ingested batch {i // BATCH_SIZE + 1}: {i} → {batch_end}")

# ---------------- FINAL CONFIRMATION ---------------- #
print("\n✅ Ingestion complete!")
print(f"Total documents stored: {collection.count()}")


'''
UPGRADES:





'''